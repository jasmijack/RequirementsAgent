import json
import csv
from docx import Document
from datetime import datetime

# Generic system configuration
system_config = {
    "platform": "Enterprise Analytics System",
    "reporting_tools": ["Tableau", "Power BI"],
    "data_sources": ["PostgreSQL", "Snowflake", "REST APIs"],
    "auth_method": "OAuth 2.0",
    "current_limitations": [
        "No automated anomaly detection",
        "Manual report distribution",
        "Siloed data sources"
    ]
}

def prompt_user_for_idea():
    print("=== Business Requirements Generator ===")
    idea = input("Enter your new feature or idea: ").strip()
    goal = input("What is the primary business goal? ").strip()
    users = input("Who are the intended users or stakeholders? ").strip()
    impact = input("What is the expected impact? ").strip()
    return idea, goal, users, impact

def create_user_story(idea, users, goal):
    return f"As a {users}, I want {idea} so that I can {goal.lower()}"

def create_technical_requirements():
    return [
        "System must support integration via REST API",
        f"Must support SSO via {system_config['auth_method']}",
        "Compatible with Tableau and Power BI visualization outputs",
        "Scheduled refresh support and error logging required"
    ]

def suggest_architecture():
    return (
        "Suggested Architecture: User > Web App > REST API > Snowflake/PostgreSQL > BI Tool (Power BI/Tableau)"
    )

def export_to_csv(idea, goal, impact):
    filename = f"Jira_Import_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
    with open(filename, mode='w', newline='') as file:
        writer = csv.DictWriter(file, fieldnames=["Summary", "Description", "Issue Type"])
        writer.writeheader()
        writer.writerow({
            "Summary": idea,
            "Description": f"Goal: {goal}\nImpact: {impact}\nPlatform: {system_config['platform']}",
            "Issue Type": "Task"
        })
    print(f"✅ Jira-ready CSV exported to {filename}")

def generate_business_requirements(idea, goal, users, impact):
    doc = Document()
    doc.add_heading("Business Requirements Document", 0)
    doc.add_paragraph(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_heading("1. Overview", level=1)
    doc.add_paragraph(f"Feature/Idea: {idea}")
    doc.add_paragraph(f"Business Goal: {goal}")

    doc.add_heading("2. Stakeholders", level=1)
    doc.add_paragraph(f"Primary Users: {users}")

    doc.add_heading("3. Expected Impact", level=1)
    doc.add_paragraph(impact)

    doc.add_heading("4. Current System Summary", level=1)
    doc.add_paragraph(f"Platform: {system_config['platform']}")
    doc.add_paragraph(f"Reporting Tools: {', '.join(system_config['reporting_tools'])}")
    doc.add_paragraph(f"Data Sources: {', '.join(system_config['data_sources'])}")
    doc.add_paragraph(f"Authentication Method: {system_config['auth_method']}")
    doc.add_paragraph("System Limitations:")
    for limitation in system_config["current_limitations"]:
        doc.add_paragraph(f"- {limitation}", style='List Bullet')

    doc.add_heading("5. User Story", level=1)
    doc.add_paragraph(create_user_story(idea, users, goal))

    doc.add_heading("6. Technical Requirements", level=1)
    for req in create_technical_requirements():
        doc.add_paragraph(f"- {req}", style='List Bullet')

    doc.add_heading("7. Suggested Architecture", level=1)
    doc.add_paragraph(suggest_architecture())

    filename = f"Business_Requirements_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(filename)
    print(f"✅ BRD exported to {filename}")


def main():
    idea, goal, users, impact = prompt_user_for_idea()
    generate_business_requirements(idea, goal, users, impact)
    export_to_csv(idea, goal, impact)

if __name__ == "__main__":
    main()
